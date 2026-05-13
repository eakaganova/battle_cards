from __future__ import annotations

import io
import json
from typing import Dict, List

import pandas as pd


STATUS_LABELS = {
    "confirmed": "подтверждено",
    "ambiguous": "неоднозначно",
    "conflicting": "конфликт",
    "missing": "нет данных",
    "inferred": "выведено",
    "needs_review": "нужна проверка",
}

INSIGHT_SECTION_LABELS = {
    "executive_summary": "Краткое резюме",
    "strengths_weaknesses": "Сильные и слабые стороны",
    "competitive_advantages": "Конкурентные преимущества",
    "gaps": "Пробелы",
    "recommendations": "Рекомендации",
    "sales_insights": "Выводы для продаж",
    "ux_insights": "UX-выводы",
    "product_conclusions": "Продуктовые выводы",
    "swot": "SWOT",
    "positioning_analysis": "Позиционирование",
    "value_proposition_comparison": "Сравнение ценностных предложений",
    "uncertainty_notes": "Неопределённость",
    "conflicts": "Конфликты данных",
}


def cells_to_dataframe(cells: Dict[str, Dict[str, object]]) -> pd.DataFrame:
    fields: List[str] = []
    for competitor_fields in cells.values():
        for field in competitor_fields.keys():
            if field not in fields:
                fields.append(field)

    rows: List[Dict[str, object]] = []
    for field in fields:
        row: Dict[str, object] = {"Параметр": field}
        for competitor, competitor_fields in cells.items():
            cell = competitor_fields.get(field)
            if not cell:
                row[competitor] = ""
                continue
            data = cell.to_dict() if hasattr(cell, "to_dict") else cell
            row[competitor] = data.get("normalized_value") or data.get("extracted_value") or ""
        rows.append(row)
    return pd.DataFrame(rows)


def cells_to_evidence_dataframe(cells: Dict[str, Dict[str, object]]) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    for competitor, fields in cells.items():
        for field, cell in fields.items():
            data = cell.to_dict() if hasattr(cell, "to_dict") else cell
            row = {
                "Конкурент": competitor,
                "Параметр": field,
                "Извлечено": data.get("extracted_value", ""),
                "Нормализовано": data.get("normalized_value", ""),
                "Статус": STATUS_LABELS.get(str(data.get("status", "")), data.get("status", "")),
                "Уверенность": data.get("confidence_score", 0),
                "Источник": data.get("source_url", ""),
                "Фрагмент": data.get("source_fragment", ""),
                "Метод": data.get("extraction_method", ""),
                "Обоснование": data.get("reasoning", ""),
            }
            rows.append(row)
    return pd.DataFrame(rows)


def export_csv(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8-sig")


def export_excel(
    df: pd.DataFrame,
    insights: Dict[str, object],
    logs: List[Dict[str, str]],
    diff: List[Dict[str, object]],
    evidence_df: pd.DataFrame | None = None,
) -> bytes:
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Таблица")
        if evidence_df is not None and not evidence_df.empty:
            evidence_df.to_excel(writer, index=False, sheet_name="Источники")
        pd.DataFrame(flatten_insights(insights)).to_excel(writer, index=False, sheet_name="Выводы")
        pd.DataFrame(logs).to_excel(writer, index=False, sheet_name="Логи")
        pd.DataFrame(diff).to_excel(writer, index=False, sheet_name="Изменения")
    return output.getvalue()


def export_markdown(df: pd.DataFrame, insights: Dict[str, object], diff: List[Dict[str, object]]) -> bytes:
    parts = ["# Конкурентное исследование", "", "## Сравнительная таблица", df.to_markdown(index=False), "", "## Выводы"]
    for key, value in insights.items():
        parts.append(f"### {INSIGHT_SECTION_LABELS.get(key, key)}")
        parts.append(json.dumps(value, ensure_ascii=False, indent=2))
    if diff:
        parts.extend(["", "## Изменения", pd.DataFrame(diff).to_markdown(index=False)])
    return "\n".join(parts).encode("utf-8")


def export_docx(df: pd.DataFrame, insights: Dict[str, object], diff: List[Dict[str, object]]) -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("Конкурентное исследование", 0)
    document.add_heading("Сравнительная таблица", level=1)
    table = document.add_table(rows=1, cols=len(df.columns))
    for index, column in enumerate(df.columns):
        table.rows[0].cells[index].text = str(column)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for index, column in enumerate(df.columns):
            cells[index].text = str(row[column])
    document.add_heading("Выводы", level=1)
    for key, value in insights.items():
        document.add_heading(INSIGHT_SECTION_LABELS.get(key, key), level=2)
        document.add_paragraph(json.dumps(value, ensure_ascii=False, indent=2))
    if diff:
        document.add_heading("Изменения", level=1)
        document.add_paragraph(json.dumps(diff, ensure_ascii=False, indent=2))
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def export_pdf(df: pd.DataFrame, insights: Dict[str, object], diff: List[Dict[str, object]]) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
    except Exception:
        return export_markdown(df, insights, diff)
    output = io.BytesIO()
    pdf = canvas.Canvas(output, pagesize=A4)
    width, height = A4
    y = height - 40
    for line in export_markdown(df.head(30), insights, diff).decode("utf-8").splitlines():
        if y < 40:
            pdf.showPage()
            y = height - 40
        pdf.drawString(35, y, line[:120])
        y -= 14
    pdf.save()
    return output.getvalue()


def google_sheets_payload(df: pd.DataFrame) -> str:
    return df.to_csv(index=False)


def flatten_insights(insights: Dict[str, object]) -> List[Dict[str, str]]:
    rows: List[Dict[str, str]] = []
    for key, value in insights.items():
        rows.append({"раздел": INSIGHT_SECTION_LABELS.get(key, key), "содержание": json.dumps(value, ensure_ascii=False)})
    return rows
